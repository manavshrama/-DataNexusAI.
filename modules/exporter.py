import pandas as pd
import io
import json
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from datetime import datetime

class ExporterModule:
    """Module for exporting data in 8 different formats."""
    
    @staticmethod
    def to_csv(df):
        return df.to_csv(index=False).encode('utf-8')

    @staticmethod
    def to_excel(df):
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Data')
            df.describe().to_excel(writer, sheet_name='Statistics')
        return output.getvalue()

    @staticmethod
    def to_json(df, orient='records'):
        return df.to_json(orient=orient, indent=2).encode('utf-8')

    @staticmethod
    def to_pdf(df, filename="report.pdf"):
        buffer = io.BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)
        p.setFont("Helvetica-Bold", 16)
        p.drawString(100, 750, "Data Nexus AI - Data Report")
        p.setFont("Helvetica", 10)
        p.drawString(100, 730, f"Generated at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        p.drawString(100, 700, f"Rows: {df.shape[0]}")
        p.drawString(100, 685, f"Columns: {df.shape[1]}")
        
        y_pos = 650
        p.drawString(100, y_pos, "Column Names:")
        y_pos -= 15
        for col in df.columns[:20]: # Limit to first 20 columns
            p.drawString(120, y_pos, f"- {col}")
            y_pos -= 15
            if y_pos < 50:
                p.showPage()
                y_pos = 750
        
        p.save()
        return buffer.getvalue()

    @staticmethod
    def to_word(df):
        doc = Document()
        doc.add_heading('Data Nexus AI - Data Report', 0)
        doc.add_paragraph(f"Report generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        doc.add_heading('Data Summary', level=1)
        doc.add_paragraph(f"Shape: {df.shape[0]} rows x {df.shape[1]} columns")
        
        doc.add_heading('Sample Data', level=1)
        table = doc.add_table(rows=1, cols=min(len(df.columns), 10))
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns[:10]):
            hdr_cells[i].text = col
            
        for i in range(min(len(df), 5)):
            row_cells = table.add_row().cells
            for j, val in enumerate(df.iloc[i][:10]):
                row_cells[j].text = str(val)
                
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()

    @staticmethod
    def to_sql(df, table_name="data_nexus"):
        create_stmt = f"CREATE TABLE {table_name} (\n"
        for col, dtype in df.dtypes.items():
            sql_type = "TEXT"
            if "int" in str(dtype) or "float" in str(dtype):
                sql_type = "NUMERIC"
            create_stmt += f"  {col} {sql_type},\n"
        create_stmt = create_stmt.rstrip(",\n") + "\n);"
        
        # Simplified INSERT generation
        inserts = ""
        for _, row in df.head(10).iterrows():
            vals = [f"'{str(v)}'" if isinstance(v, str) else str(v) for v in row.values]
            inserts += f"INSERT INTO {table_name} VALUES ({', '.join(vals)});\n"
            
        return (create_stmt + "\n\n" + inserts).encode('utf-8')

    @staticmethod
    def to_markdown(df):
        return df.head(100).to_markdown().encode('utf-8')

    @staticmethod
    def to_html(df):
        return df.head(100).to_html(classes='table table-striped table-hover').encode('utf-8')
